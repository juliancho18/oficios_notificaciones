import os
import re
import unicodedata
from pathlib import Path
from typing import Dict, Tuple, Optional, List, Iterable, Any

import pandas as pd


def find_col(df: pd.DataFrame, posibles: List[str]) -> Optional[str]:
    def norm(s):
        s = "" if s is None else str(s)
        s = s.strip().lower()
        s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii")
        s = s.replace(" ", "").replace("_", "").replace(".", "").replace(":", "").replace("-", "")
        return s

    cols = list(df.columns)
    cols_norm = {norm(c): c for c in cols}

    for p in posibles:
        pn = norm(p)
        if pn in cols_norm:
            return cols_norm[pn]

    for p in posibles:
        pn = norm(p)
        for c in cols:
            if pn in norm(c):
                return c

    return None


def debug_detectar_columnas(df_filtrado: pd.DataFrame) -> Dict[str, Optional[str]]:
    print("Columnas DF (top 30):")
    print(list(df_filtrado.columns)[:30])

    col_notificador = find_col(df_filtrado, ["Notificador/Revisor", "Notificador", "Revisor"])
    col_estado = find_col(
        df_filtrado,
        [
            "Estado_Actuacion",
            "Estado Actuacion",
            "Estado",
            "Estado_Actuación",
            "Estado Actuación",
            "Estado de la Actuacion",
            "Estado de la Actuación",
        ],
    )
    col_tipo_oficio = find_col(
        df_filtrado,
        ["Tipo_Oficio", "Tipo Oficio", "Tipo de oficio", "Tipo Oficio a generar", "ACTUACION"],
    )
    col_expediente = find_col(
        df_filtrado,
        ["numero_expediente", "No. Exp.", "No Exp", "Expediente", "No_Expediente"],
    )
    col_notificado = find_col(
        df_filtrado,
        ["Nombre_Notificado", "Nombre Notificado", "Notificado", "Nombre del Notificado"],
    )

    print("\nDetectadas:")
    print("col_notificador:", col_notificador)
    print("col_estado     :", col_estado)
    print("col_tipo_oficio:", col_tipo_oficio)
    print("col_expediente :", col_expediente)
    print("col_notificado :", col_notificado)

    faltan = [
        k
        for k, v in {
            "col_notificador": col_notificador,
            "col_estado": col_estado,
            "col_tipo_oficio": col_tipo_oficio,
            "col_expediente": col_expediente,
            "col_notificado": col_notificado,
        }.items()
        if v is None
    ]

    if faltan:
        raise KeyError(f"❌ No pude detectar estas columnas: {faltan}. Revisa nombres en df_filtrado.columns.")

    return {
        "col_notificador": col_notificador,
        "col_estado": col_estado,
        "col_tipo_oficio": col_tipo_oficio,
        "col_expediente": col_expediente,
        "col_notificado": col_notificado,
    }


class Utils:
    # =========================
    # Paths y compatibilidad con notebooks previos
    # =========================
    def ruta(self, *partes: str) -> str:
        return str(Path(*partes))

    def asegurar_carpeta(self, ruta: str) -> str:
        Path(ruta).mkdir(parents=True, exist_ok=True)
        return ruta

    def _ensure_dir(self, path: str) -> str:
        os.makedirs(path, exist_ok=True)
        return path

    def _ruta_base_proyecto(self) -> Path:
        return Path(__file__).resolve().parent

    def _ruta_plantillas(self) -> Path:
        ruta = self._ruta_base_proyecto() / "plantillas"
        ruta.mkdir(exist_ok=True)
        return ruta

    # =========================
    # Helpers de limpieza
    # =========================
    def normalizar_texto(self, s: str) -> str:
        if s is None:
            return ""
        s = str(s).strip()
        s = re.sub(r"\s+", " ", s)
        return s

    def _norm_txt(self, x: str) -> str:
        x = "" if x is None else str(x)
        x = x.strip().upper()
        x = " ".join(x.split())
        x = unicodedata.normalize("NFKD", x).encode("ascii", "ignore").decode("ascii")
        return x

    def quitar_tildes(self, texto: str) -> str:
        texto = "" if texto is None else str(texto)
        return unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")

    def normalizar_nombre_persona(self, nombre: str) -> str:
        nombre = "" if nombre is None else str(nombre)
        nombre = self.quitar_tildes(nombre).upper().strip()
        nombre = re.sub(r"\s+", " ", nombre)
        return nombre

    def limpiar_nombre_archivo(self, s: str, max_len: int = 120) -> str:
        s = self.normalizar_texto(s)
        if s == "":
            s = "SIN_NOMBRE"

        s = re.sub(r'[\\/*?:"<>|]', "", s)
        s = "".join(ch for ch in s if unicodedata.category(ch)[0] != "C")
        s = s.replace(" ", "_")
        s = s[:max_len].rstrip("_")
        return s or "SIN_NOMBRE"

    def _safe(self, s: str, max_len: int = 120) -> str:
        return self.limpiar_nombre_archivo(s, max_len=max_len)

    def limpiar_expediente(self, valor) -> str:
        if pd.isna(valor):
            return ""
        txt = str(valor).strip()
        txt = re.sub(r"\.0$", "", txt)
        txt = re.sub(r"\s+", "", txt)
        txt = txt.replace("\n", "").replace("\r", "")
        return txt

    # =========================
    # Normalizadores de columnas
    # =========================
    def estandarizar_columnas(self, df: pd.DataFrame) -> pd.DataFrame:
        df_out = df.copy()
        nuevas = []
        for col in df_out.columns:
            txt = "" if col is None else str(col)
            txt = txt.strip()
            txt = unicodedata.normalize("NFKD", txt).encode("ascii", "ignore").decode("ascii")
            txt = txt.lower()
            txt = re.sub(r"[^a-z0-9]+", "_", txt)
            txt = re.sub(r"_+", "_", txt).strip("_")
            nuevas.append(txt if txt else "columna")
        df_out.columns = self._hacer_columnas_unicas_basico(nuevas)
        return df_out

    def limpiar_filas_vacias(self, df: pd.DataFrame) -> pd.DataFrame:
        return df.dropna(how="all").reset_index(drop=True)

    def hacer_columnas_unicas(self, columnas: Iterable[Any]) -> List[str]:
        return self._hacer_columnas_unicas_basico(list(columnas))

    # =========================
    # del / de la
    # =========================
    def _norm_lower_sin_tildes(self, x: str) -> str:
        x = "" if x is None else str(x)
        x = x.strip().lower()
        x = " ".join(x.split())
        x = unicodedata.normalize("NFKD", x).encode("ascii", "ignore").decode("ascii")
        return x

    def obtener_de_la_o_del(self, valor: str) -> str:
        texto = self._norm_lower_sin_tildes(valor)
        if texto == "":
            return "de la"

        primera = texto.split()[0]

        femeninas = {
            "resolucion",
            "actuacion",
            "notificacion",
            "comunicacion",
            "liquidacion",
            "sancion",
            "certificacion",
        }
        masculinas = {
            "auto",
            "informe",
            "oficio",
            "acto",
            "memorando",
            "concepto",
        }

        if primera in femeninas:
            return "de la"
        if primera in masculinas:
            return "del"
        if primera.endswith("a"):
            return "de la"
        return "del"

    def agregar_columnas_del_de_la(
        self,
        df: pd.DataFrame,
        col_origen: str,
        col_articulo: str = "articulo_preposicion",
        col_frase: str = "descripcion_con_articulo",
        minusculas: bool = True,
    ) -> pd.DataFrame:
        if col_origen not in df.columns:
            raise KeyError(
                f"No existe la columna '{col_origen}' en el DataFrame.\n"
                f"Columnas disponibles: {df.columns.tolist()}"
            )

        out = df.copy()
        out[col_articulo] = out[col_origen].apply(self.obtener_de_la_o_del)
        base = out[col_origen].astype(str).str.strip().str.lower() if minusculas else out[col_origen].astype(str).str.strip()
        out[col_frase] = out[col_articulo] + " " + base
        return out

    # =========================
    # Funciones base
    # =========================
    def filtrar_columnas(self, df: pd.DataFrame, vector: list) -> pd.DataFrame:
        cols_existentes = [c for c in vector if c in df.columns]
        faltantes = [c for c in vector if c not in df.columns]
        if faltantes:
            print(f"⚠️ Columnas no encontradas y omitidas: {faltantes}")
        return df[cols_existentes].copy()

    def fusionar_dataframes(self, df1: pd.DataFrame, df2: pd.DataFrame, on, how: str = "left") -> pd.DataFrame:
        return df1.merge(df2, on=on, how=how)

    def agregar_dummies(self, df: pd.DataFrame, col: str) -> pd.DataFrame:
        df_out = df.copy()
        dummies = pd.get_dummies(df_out[col], prefix=col, dummy_na=False)
        return pd.concat([df_out.drop(columns=[col]), dummies], axis=1)

    def generar_features_fecha(self, df: pd.DataFrame, col_fecha: str) -> pd.DataFrame:
        df_out = df.copy()
        df_out[col_fecha] = pd.to_datetime(df_out[col_fecha], errors="coerce")
        df_out[f"{col_fecha}_anio"] = df_out[col_fecha].dt.year
        df_out[f"{col_fecha}_mes"] = df_out[col_fecha].dt.month
        df_out[f"{col_fecha}_dia"] = df_out[col_fecha].dt.day
        df_out[f"{col_fecha}_dia_semana"] = df_out[col_fecha].dt.day_name()
        return df_out

    def agregar_columna_binaria(self, df: pd.DataFrame, col: str, valor_objetivo, nueva_columna: str) -> pd.DataFrame:
        df_out = df.copy()
        df_out[nueva_columna] = (df_out[col] == valor_objetivo).astype(int)
        return df_out

    def asignar_grupos(self, df: pd.DataFrame, col: str) -> pd.DataFrame:
        df_out = df.copy()
        grupo1 = {"jeniffer caballero", "cristian gil", "valentina bernal"}
        df_out["grupo"] = (
            df_out[col]
            .astype(str)
            .str.strip()
            .str.lower()
            .apply(lambda x: 1 if x in grupo1 else 2)
        )
        return df_out

    # =========================
    # Regional -> indicador, director, dirección, municipio
    # =========================
    def agregar_info_regional(self, df: pd.DataFrame, col_regional: str = "Regional") -> pd.DataFrame:
        if col_regional not in df.columns:
            raise KeyError(
                f"No existe la columna '{col_regional}' en el DataFrame.\n"
                f"Columnas disponibles: {df.columns.tolist()}"
            )

        regional_map: Dict[str, Tuple[str, str, str, str]] = {
            self._norm_txt("ALMEIDAS Y GUATAVITA"): (
                "DRAG", "HECTOR JOSE CAMACHO ACOSTA",
                "Carrera. 5 N° 5-73 piso 2 y 4 Edificio Molino del Parque",
                "Chocontá, Cundinamarca",
            ),
            self._norm_txt("ALTO MAGDALENA"): (
                "DRAM", "CAMILA ANDREA VELASQUEZ BAQUERO",
                "Calle 21 No. 8 – 23 Barrio Granada",
                "Girardot, Cundinamarca",
            ),
            self._norm_txt("BAJO MAGDALENA"): (
                "DRBM", "JUAN FILIBERTO COTRINO GUEVARA",
                "Calle 4 N° 5-68 Camellón Real",
                "Guaduas, Cundinamarca",
            ),
            self._norm_txt("CHIQUINQUIRA"): (
                "DRCH", "YIBER ESTEBAN GONZALEZ GIL",
                "Carrera 6 N° 9-40",
                "Chiquinquirá, Boyacá",
            ),
            self._norm_txt("GUALIVA"): (
                "DRGU", "RONALD DAVID PRIETO GONZALEZ",
                "Calle. 8 N° 10ª - 41 Barrio El Recreo",
                "Villeta, Cundinamarca",
            ),
            self._norm_txt("MAGDALENA CENTRO"): (
                "DRMC", "JUAN CARLOS ESCOBAR CRISTANCHO",
                "Carrera 5 N° 3 - 02 Esquina",
                "Vianí, Cundinamarca",
            ),
            self._norm_txt("RIONEGRO"): (
                "DRRN", "ALEJANDRO FUQUITVA CASALLAS",
                "Calle 9 N° 19-72, Barrio Antonio Nariño",
                "Pacho, Cundinamarca",
            ),
            self._norm_txt("SABANA CENTRO"): (
                "DRSC", "ANDRES MAURICIO GARZON ORJUELA",
                "Calle 7 A N°11- 40 Barrio Algarra",
                "Zipaquirá, Cundinamarca",
            ),
            self._norm_txt("SABANA OCCIDENTE"): (
                "DRSO", "LINA CAMILA CORTES ACOSTA",
                "Carrera 4. N° 4-38",
                "Facatativá, Cundinamarca",
            ),
            self._norm_txt("SOACHA"): (
                "DRSOA", "CESAR AUGUSTO RICO MAYORGA",
                "Trans. 7 F N° 26-38 Barrio El Nogal - Soacha",
                "Soacha, Cundinamarca",
            ),
            self._norm_txt("SUMAPAZ"): (
                "DRSU", "ERIKA ALVAREZ CASTAÑEDA",
                "Avenida las Palmas N° 15-17",
                "Fusagasugá, Cundinamarca",
            ),
            self._norm_txt("TEQUENDAMA"): (
                "DRTE", "NIDIA CRUZ ORTEGA",
                "Carrera 21 con Calle 2da, Esquina Barrio El Recreo",
                "La Mesa, Cundinamarca",
            ),
            self._norm_txt("UBATE"): (
                "DRUB", "JULIO CESAR SIERRA LEON",
                "Transversal 2da. N°1E-40",
                "Ubaté, Cundinamarca",
            ),
            self._norm_txt("BOGOTA LA CALERA"): (
                "DRBC", "YUBER YESID CARDENAS PULIDO",
                "Carrera 10 N° 16-82 piso 4",
                "Bogotá, Cundinamarca",
            ),
        }

        df_out = df.copy()
        df_out["_regional_norm"] = df_out[col_regional].map(self._norm_txt)

        def _get_tuple(k: str) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str]]:
            return regional_map.get(k, (None, None, None, None))

        df_out["INDICADOR"] = df_out["_regional_norm"].map(lambda k: _get_tuple(k)[0])
        df_out["DIRECTOR_REGIONAL"] = df_out["_regional_norm"].map(lambda k: _get_tuple(k)[1])
        df_out["DIRECCION_REGIONAL"] = df_out["_regional_norm"].map(lambda k: _get_tuple(k)[2])
        df_out["MUNICIPIO_REGIONAL"] = df_out["_regional_norm"].map(lambda k: _get_tuple(k)[3])

        faltan = df_out[
            df_out["INDICADOR"].isna()
            | df_out["DIRECTOR_REGIONAL"].isna()
            | df_out["DIRECCION_REGIONAL"].isna()
            | df_out["MUNICIPIO_REGIONAL"].isna()
        ][col_regional].astype(str).unique().tolist()

        df_out.drop(columns=["_regional_norm"], inplace=True)

        if faltan:
            raise ValueError(
                "❌ Hay Regional(es) sin mapeo completo (INDICADOR/DIRECTOR/DIRECCION/MUNICIPIO):\n"
                + "\n".join([f"- {x}" for x in faltan])
                + "\n\nSolución: agrega esas regionales al diccionario regional_map."
            )

        return df_out

    # =========================
    # Cruce de correspondencia
    # =========================
    def agregar_fila_cruce_excel(self, df: pd.DataFrame, col_salida: str = "No_Fila_Cruce", offset: int = 2) -> pd.DataFrame:
        df_out = df.copy()
        df_out[col_salida] = range(offset, offset + len(df_out))
        return df_out

    def agregar_fila_cruce_word(self, df: pd.DataFrame, col_salida: str = "No_Fila_Cruce", offset: int = 3) -> pd.DataFrame:
        df_out = df.copy()
        df_out[col_salida] = range(offset, offset + len(df_out))
        return df_out

    def agregar_fila_cruce(self, df: pd.DataFrame, col_salida: str = "No_Fila_Cruce", offset: int = 3) -> pd.DataFrame:
        return self.agregar_fila_cruce_word(df, col_salida=col_salida, offset=offset)

    # =========================
    # Exportar Excel por persona
    # =========================
    def exportar_excels_por_persona(
        self,
        df: pd.DataFrame,
        col_persona: str,
        carpeta_salida: str = "salidas_digitadores",
        nombre_prefix: str = "BASE",
        add_fila_cruce: bool = True,
        col_fila_cruce: str = "No_Fila_Cruce",
        offset_fila_cruce: int = 3,
    ) -> pd.DataFrame:
        os.makedirs(carpeta_salida, exist_ok=True)

        df_work = df.copy()
        df_work[col_persona] = df_work[col_persona].astype(str).apply(self.normalizar_texto)

        resumen = []
        for persona, sub in df_work.groupby(col_persona, dropna=False):
            persona_str = self.normalizar_texto(persona)
            sub_out = sub.copy().reset_index(drop=True)

            if add_fila_cruce:
                sub_out = self.agregar_fila_cruce(sub_out, col_salida=col_fila_cruce, offset=offset_fila_cruce)

            safe_name = self.limpiar_nombre_archivo(persona_str)
            filename = f"{nombre_prefix}_{safe_name}.xlsx"
            ruta = os.path.join(carpeta_salida, filename)

            sub_out.to_excel(ruta, index=False)
            resumen.append((persona_str, len(sub_out), ruta))

        return pd.DataFrame(resumen, columns=["persona", "filas", "ruta_archivo"])

    # =========================
    # Partir DF por persona + cruce
    # =========================
    def partir_por_persona_con_cruce(
        self,
        df: pd.DataFrame,
        col_persona: str = "Notificador/Revisor",
        col_cruce: str = "No_Fila_Cruce",
        offset_cruce: int = 3,
    ) -> Dict[str, pd.DataFrame]:
        if col_persona not in df.columns:
            raise KeyError(
                f"No existe la columna '{col_persona}' en el DataFrame.\n"
                f"Columnas disponibles: {df.columns.tolist()}"
            )

        df_work = df.copy()
        df_work[col_persona] = df_work[col_persona].astype(str).apply(self.normalizar_texto)

        salida: Dict[str, pd.DataFrame] = {}
        for persona, sub in df_work.groupby(col_persona, dropna=False):
            sub_out = sub.copy().reset_index(drop=True)
            sub_out = self.agregar_fila_cruce(sub_out, col_salida=col_cruce, offset=offset_cruce)
            salida[str(persona)] = sub_out

        return salida

    # =========================
    # Lógica: ejecutorias / tipo oficio
    # =========================
    def _contiene_ejecutoria(self, x: str) -> bool:
        if x is None:
            return False
        txt = str(x).strip().lower()
        txt = unicodedata.normalize("NFKD", txt).encode("ascii", "ignore").decode("ascii")
        txt = " ".join(txt.split())
        patrones = ["ejecutoria", "para ejecutoria", "en ejecutoria", "ejecutoriado"]
        return any(p in txt for p in patrones)

    def _norm_tipo_oficio(self, x: str) -> str:
        t = "" if x is None else str(x).strip().lower()
        t = " ".join(t.split())
        t = unicodedata.normalize("NFKD", t).encode("ascii", "ignore").decode("ascii")
        if "cita" in t:
            return "citacion"
        if "comun" in t:
            return "comunicacion"
        if "electr" in t or "correo" in t:
            return "notificacion_electronica"
        if "aviso" in t:
            return "aviso"
        return t if t else "sin_tipo"

    def resolver_columnas_operativas(self, df: pd.DataFrame) -> Dict[str, str]:
        detectadas = debug_detectar_columnas(df)
        return {
            "col_notificador": detectadas["col_notificador"],
            "col_estado": detectadas["col_estado"],
            "col_tipo_oficio": detectadas["col_tipo_oficio"],
            "col_expediente": detectadas["col_expediente"],
            "col_notificado": detectadas["col_notificado"],
        }

    # =========================
    # Excel operativo por abogado
    # =========================
    def exportar_excel_oficios_y_ejecutorias_por_notificador(
        self,
        df: pd.DataFrame,
        col_notificador: str = "Notificador/Revisor",
        col_estado_actuacion: str = "Estado_Actuacion",
        col_tipo_oficio: str = "Tipo_Oficio",
        col_expediente: str = "numero_expediente",
        col_notificado: str = "Nombre_Notificado",
        carpeta_salida: str = "output_abogados",
        nombre_archivo: str = "RESUMEN_OFICIOS_Y_EJECUTORIAS.xlsx",
        add_fila_cruce: bool = True,
        col_fila_cruce: str = "No_Fila_Cruce",
        offset_fila_cruce: int = 3,
        autodetectar_columnas: bool = True,
    ) -> pd.DataFrame:
        if autodetectar_columnas:
            detectadas = self.resolver_columnas_operativas(df)
            col_notificador = detectadas["col_notificador"]
            col_estado_actuacion = detectadas["col_estado"]
            col_tipo_oficio = detectadas["col_tipo_oficio"]
            col_expediente = detectadas["col_expediente"]
            col_notificado = detectadas["col_notificado"]

        for c in [col_notificador, col_estado_actuacion, col_tipo_oficio, col_expediente, col_notificado]:
            if c not in df.columns:
                raise KeyError(f"No existe la columna '{c}'. Columnas disponibles: {df.columns.tolist()}")

        self._ensure_dir(carpeta_salida)
        dfw = df.copy()
        dfw[col_notificador] = dfw[col_notificador].astype(str).apply(self.normalizar_texto)

        resumen = []
        for abogado, sub in dfw.groupby(col_notificador, dropna=False):
            abogado_txt = self.normalizar_texto(abogado)
            carpeta_abogado = self._ensure_dir(os.path.join(carpeta_salida, self._safe(abogado_txt)))

            mask_ej = sub[col_estado_actuacion].apply(self._contiene_ejecutoria) | sub[col_tipo_oficio].apply(self._contiene_ejecutoria)
            df_ej = sub[mask_ej].copy().reset_index(drop=True)
            df_of = sub[~mask_ej].copy().reset_index(drop=True)

            df_of_tab = df_of[[col_tipo_oficio, col_expediente, col_notificado, col_estado_actuacion]].copy()
            df_of_tab.rename(
                columns={
                    col_tipo_oficio: "tipo_oficio",
                    col_expediente: "numero_expediente",
                    col_notificado: "notificado",
                    col_estado_actuacion: "estado_actuacion",
                },
                inplace=True,
            )

            df_ej_tab = df_ej[[col_expediente, col_notificado, col_estado_actuacion]].copy()
            df_ej_tab.rename(
                columns={
                    col_expediente: "numero_expediente",
                    col_notificado: "notificado",
                    col_estado_actuacion: "estado_actuacion",
                },
                inplace=True,
            )

            if add_fila_cruce:
                df_of_tab = self.agregar_fila_cruce(df_of_tab, col_salida=col_fila_cruce, offset=offset_fila_cruce)
                df_ej_tab = self.agregar_fila_cruce(df_ej_tab, col_salida=col_fila_cruce, offset=offset_fila_cruce)

            ruta = os.path.join(carpeta_abogado, nombre_archivo)
            with pd.ExcelWriter(ruta, engine="openpyxl") as writer:
                df_of_tab.to_excel(writer, index=False, sheet_name="OFICIOS_POR_CREAR")
                df_ej_tab.to_excel(writer, index=False, sheet_name="EJECUTORIAS")

            resumen.append(
                {
                    "abogado": abogado_txt,
                    "oficios_por_crear": len(df_of_tab),
                    "ejecutorias": len(df_ej_tab),
                    "ruta_excel": ruta,
                }
            )

        return pd.DataFrame(resumen)

    # =========================
    # Helpers Word desde plantillas externas
    # =========================
    def _mapa_plantillas_docx(self) -> Dict[str, str]:
        return {
            "aviso": "Formato Aviso Recurso (AUTOMATIZADO).docx",
            "citacion": "Formato Citación CPACA2011 (Automatizado).docx",
            "comunicacion": "Formato Comunicación (Automatizado).docx",
            "notificacion_electronica": "Notificación electrónica (Automatizado).docx",
        }

    def _obtener_ruta_plantilla_por_tipo(self, tipo_oficio: str) -> str:
        tipo = self._norm_tipo_oficio(tipo_oficio)
        mapa = self._mapa_plantillas_docx()
        if tipo not in mapa:
            raise ValueError(f"Tipo de oficio no reconocido: {tipo_oficio}. Tipos válidos: {list(mapa.keys())}")
        ruta = self._ruta_plantillas() / mapa[tipo]
        if not ruta.exists():
            raise FileNotFoundError(f"No existe la plantilla: {ruta}")
        return str(ruta)

    def _reemplazar_texto_en_parrafo(self, parrafo, replacements: Dict[str, str]) -> None:
        texto_original = parrafo.text
        texto_nuevo = texto_original
        for buscar, reemplazar in replacements.items():
            if buscar in texto_nuevo:
                texto_nuevo = texto_nuevo.replace(buscar, reemplazar)
        if texto_nuevo == texto_original:
            return

        estilo = parrafo.runs[0] if parrafo.runs else None
        for run in parrafo.runs[::-1]:
            parrafo._element.remove(run._element)
        nuevo_run = parrafo.add_run(texto_nuevo)

        if estilo:
            nuevo_run.bold = estilo.bold
            nuevo_run.italic = estilo.italic
            nuevo_run.underline = estilo.underline
            nuevo_run.font.name = estilo.font.name
            nuevo_run.font.size = estilo.font.size

    def _replace_in_doc_generico(self, doc, replacements: Dict[str, str]) -> None:
        for p in doc.paragraphs:
            self._reemplazar_texto_en_parrafo(p, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._reemplazar_texto_en_parrafo(p, replacements)
        for section in doc.sections:
            for p in section.header.paragraphs:
                self._reemplazar_texto_en_parrafo(p, replacements)
            for p in section.footer.paragraphs:
                self._reemplazar_texto_en_parrafo(p, replacements)
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._reemplazar_texto_en_parrafo(p, replacements)
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            self._reemplazar_texto_en_parrafo(p, replacements)

    def _armar_replacements_desde_row(
        self,
        row: pd.Series,
        columnas_df: List[str],
        col_map: Optional[Dict[str, str]] = None,
    ) -> Dict[str, str]:
        if col_map is None:
            col_map = {}

        def _norm_key(x: str) -> str:
            x = "" if x is None else str(x)
            x = x.strip().lower()
            x = unicodedata.normalize("NFKD", x).encode("ascii", "ignore").decode("ascii")
            x = re.sub(r"[^a-z0-9]+", "", x)
            return x

        def _val(col_name: str) -> str:
            if col_name in row.index:
                v = row.get(col_name, "")
                if pd.isna(v):
                    return ""
                return str(v)
            return ""

        replacements: Dict[str, str] = {}
        for col in columnas_df:
            valor = _val(col)
            variantes = {str(col), str(col).strip(), _norm_key(col)}
            for k in variantes:
                replacements[f"«{k}»"] = valor
                replacements[f"<<{k}>>"] = valor
                replacements[f"[{k}]"] = valor
                replacements[f"{{{{{k}}}}}"] = valor

        for placeholder, real_col in col_map.items():
            valor = _val(real_col)
            variantes = {str(placeholder), str(placeholder).strip(), _norm_key(placeholder)}
            for k in variantes:
                replacements[f"«{k}»"] = valor
                replacements[f"<<{k}>>"] = valor
                replacements[f"[{k}]"] = valor
                replacements[f"{{{{{k}}}}}"] = valor

        return replacements

    # =========================
    # Generar .docx desde plantilla única
    # =========================
    def generar_words_desde_plantilla(
        self,
        df: pd.DataFrame,
        plantilla_path: str,
        carpeta_salida: str = "output_abogados",
        col_notificador: str = "Notificador/Revisor",
        col_subcarpeta: str = "ACTUACION",
        nombre_archivo_cols: Optional[List[str]] = None,
        col_map: Optional[Dict[str, str]] = None,
    ) -> pd.DataFrame:
        try:
            from docx import Document
        except Exception as e:
            raise ImportError("No está instalada la librería python-docx. Instala con:\npip install python-docx\n" f"Detalle: {e}")

        if nombre_archivo_cols is None:
            nombre_archivo_cols = ["numero_expediente"]
        if col_map is None:
            col_map = {}
        if not os.path.exists(plantilla_path):
            raise FileNotFoundError(f"No existe la plantilla: {plantilla_path}")

        self._ensure_dir(carpeta_salida)
        generados = []

        for i, row in df.iterrows():
            abogado = self._safe(row.get(col_notificador, "SIN_ABOGADO"))
            sub = self._safe(row.get(col_subcarpeta, "SIN_TIPO")) if col_subcarpeta in df.columns else "SIN_TIPO"
            out_dir = os.path.join(carpeta_salida, abogado, "WORDS", sub)
            self._ensure_dir(out_dir)

            replacements = self._armar_replacements_desde_row(row=row, columnas_df=list(df.columns), col_map=col_map)
            partes = []
            for c in nombre_archivo_cols:
                if c in df.columns:
                    partes.append(self._safe(row.get(c, "")))
            if not partes:
                partes = [f"documento_{i+1}"]

            nombre = "OFICIO_" + "_".join([p for p in partes if p]) + ".docx"
            ruta = os.path.join(out_dir, nombre)

            doc = Document(plantilla_path)
            self._replace_in_doc_generico(doc, replacements)
            doc.save(ruta)

            generados.append({"fila": i, "ruta_docx": ruta, "abogado": abogado, "subcarpeta": sub, "archivo": nombre})

        return pd.DataFrame(generados)

    # =========================
    # Generar .docx según tipo de oficio
    # =========================
    def generar_words_desde_tipo_oficio(
        self,
        df: pd.DataFrame,
        carpeta_salida: str = "output_abogados",
        col_notificador: str = "Notificador/Revisor",
        col_estado_actuacion: str = "Estado_Actuacion",
        col_tipo_oficio: str = "Tipo_Oficio",
        col_subcarpeta: Optional[str] = None,
        nombre_archivo_cols: Optional[List[str]] = None,
        col_map: Optional[Dict[str, str]] = None,
        autodetectar_columnas: bool = True,
    ) -> pd.DataFrame:
        try:
            from docx import Document
        except Exception as e:
            raise ImportError("No está instalada la librería python-docx. Instala con:\npip install python-docx\n" f"Detalle: {e}")

        if nombre_archivo_cols is None:
            nombre_archivo_cols = ["numero_expediente", "Nombre_Notificado"]
        if col_map is None:
            col_map = {}

        if autodetectar_columnas:
            detectadas = self.resolver_columnas_operativas(df)
            col_notificador = detectadas["col_notificador"]
            col_estado_actuacion = detectadas["col_estado"]
            col_tipo_oficio = detectadas["col_tipo_oficio"]

        for c in [col_notificador, col_estado_actuacion, col_tipo_oficio]:
            if c not in df.columns:
                raise KeyError(f"No existe la columna '{c}'. Columnas disponibles: {df.columns.tolist()}")

        generados = []
        dfw = df.copy()
        mask_ej_estado = dfw[col_estado_actuacion].apply(self._contiene_ejecutoria)
        mask_ej_tipo = dfw[col_tipo_oficio].apply(self._contiene_ejecutoria)
        dfw = dfw[~(mask_ej_estado | mask_ej_tipo)].copy()

        for i, row in dfw.reset_index(drop=True).iterrows():
            abogado = self.normalizar_texto(row.get(col_notificador, "SIN_ABOGADO"))
            tipo_original = row.get(col_tipo_oficio, "")
            tipo = self._norm_tipo_oficio(tipo_original)

            if tipo not in {"aviso", "citacion", "comunicacion", "notificacion_electronica"}:
                generados.append({"fila": i, "abogado": abogado, "tipo_oficio": tipo, "error": f"Tipo no válido para generar entregable: {tipo_original}"})
                continue

            try:
                plantilla_path = self._obtener_ruta_plantilla_por_tipo(tipo)
                sub = self._safe(tipo if col_subcarpeta is None else row.get(col_subcarpeta, tipo))
                out_dir = os.path.join(carpeta_salida, self._safe(abogado), "WORDS", sub)
                self._ensure_dir(out_dir)

                replacements = self._armar_replacements_desde_row(row=row, columnas_df=list(dfw.columns), col_map=col_map)
                partes = []
                for c in nombre_archivo_cols:
                    if c in dfw.columns:
                        partes.append(self._safe(row.get(c, "")))
                if not partes:
                    partes = [f"documento_{i+1}"]

                nombre = "OFICIO_" + "_".join([p for p in partes if p]) + ".docx"
                ruta = os.path.join(out_dir, nombre)

                doc = Document(plantilla_path)
                self._replace_in_doc_generico(doc, replacements)
                doc.save(ruta)

                generados.append({
                    "fila": i,
                    "abogado": abogado,
                    "tipo_oficio": tipo,
                    "plantilla": os.path.basename(plantilla_path),
                    "ruta_docx": ruta,
                    "archivo": nombre,
                })
            except Exception as e:
                generados.append({"fila": i, "abogado": abogado, "tipo_oficio": tipo, "error": str(e)})

        return pd.DataFrame(generados)

    # =========================
    # Encabezados SIDCAR
    # =========================
    def _texto_encabezado(self, valor):
        if pd.isna(valor):
            return ""
        texto = str(valor).strip()
        texto = re.sub(r"\s+", " ", texto)
        if texto.lower().startswith("unnamed:"):
            return ""
        return texto

    def _unir_dos_filas_encabezado(self, fila_1, fila_2):
        encabezados = []
        for a, b in zip(fila_1, fila_2):
            a = self._texto_encabezado(a)
            b = self._texto_encabezado(b)
            if a and b:
                encabezado = a if a.lower() == b.lower() else f"{a} {b}"
            elif a:
                encabezado = a
            elif b:
                encabezado = b
            else:
                encabezado = "columna"
            encabezado = re.sub(r"\s+", " ", encabezado).strip()
            encabezados.append(encabezado)
        return encabezados

    def _hacer_columnas_unicas_basico(self, columnas):
        conteo = {}
        resultado = []
        for col in columnas:
            base = str(col).strip() if str(col).strip() else "columna"
            if base not in conteo:
                conteo[base] = 0
                resultado.append(base)
            else:
                conteo[base] += 1
                resultado.append(f"{base}_{conteo[base]}")
        return resultado

    def normalizar_encabezados_sidcar(self, df: pd.DataFrame) -> pd.DataFrame:
        df_out = df.copy()
        columnas = []
        for col in df_out.columns:
            txt = self._texto_encabezado(col)
            txt = re.sub(r"\s+", " ", txt).strip()
            columnas.append(txt if txt else "columna")
        df_out.columns = self._hacer_columnas_unicas_basico(columnas)
        return df_out

    def cargar_excel_con_encabezado_doble(self, ruta_archivo, filas_encabezado=2):
        extension = str(ruta_archivo).lower()
        if extension.endswith(".xls"):
            df_raw = pd.read_excel(ruta_archivo, header=None, engine="xlrd")
        else:
            df_raw = pd.read_excel(ruta_archivo, header=None)

        if df_raw.shape[0] < filas_encabezado:
            raise ValueError(f"El archivo {ruta_archivo} no tiene suficientes filas para construir encabezado.")

        fila_1 = df_raw.iloc[0].tolist()
        fila_2 = df_raw.iloc[1].tolist()
        columnas_limpias = self._unir_dos_filas_encabezado(fila_1, fila_2)
        columnas_limpias = self._hacer_columnas_unicas_basico(columnas_limpias)

        df = df_raw.iloc[filas_encabezado:].copy()
        df.columns = columnas_limpias
        df = df.reset_index(drop=True)
        df = df.dropna(how="all").reset_index(drop=True)
        return df

    # =========================
    # SIDCAR - GRUPO A
    # Solo PRERADICADOS + RADICADOS
    # =========================
    def homologar_nombre_grupo(self, nombre: str, nombres_canonicos: list, cutoff: float = 0.80) -> str:
        import difflib
        nombre_norm = self.normalizar_nombre_persona(nombre)
        canonicos_norm = [self.normalizar_nombre_persona(x) for x in nombres_canonicos]
        if nombre_norm in canonicos_norm:
            return nombre_norm
        match = difflib.get_close_matches(nombre_norm, canonicos_norm, n=1, cutoff=cutoff)
        if match:
            return match[0]
        return nombre_norm

    def extraer_expediente_texto(self, texto) -> str:
        txt = "" if pd.isna(texto) else str(texto)
        txt = txt.strip()
        if not txt:
            return ""
        m = re.search(r"(?i)expediente\s*[:#\-–]?\s*([A-Za-z0-9\-\/]{4,})", txt)
        if m:
            return self.limpiar_expediente(m.group(1))
        nums = re.findall(r"\b[A-Za-z0-9\-\/]{4,}\b", txt)
        if nums:
            return self.limpiar_expediente(nums[-1])
        return ""

    def limpiar_expediente_texto(self, valor) -> str:
        return self.limpiar_expediente(valor)

    def extraer_expediente_desde_texto(self, texto: str) -> Optional[str]:
        if texto is None or pd.isna(texto):
            return None
        texto = str(texto)
        patrones = [
            r"expediente\s*[:#\-]?\s*([A-Za-z0-9\-\/]+)",
            r"exp\s*[:#\-]?\s*([A-Za-z0-9\-\/]+)",
        ]
        for patron in patrones:
            m = re.search(patron, texto, flags=re.IGNORECASE)
            if m:
                return self.limpiar_expediente_texto(m.group(1))
        return None

    def detectar_columna_persona_por_contenido(self, df: pd.DataFrame, nombres_grupo: list, top_n: int = 200) -> Optional[str]:
        grupo_norm = [self.normalizar_nombre_persona(x) for x in nombres_grupo]
        mejor_col = None
        mejor_score = -1

        for col in df.columns:
            serie = df[col].dropna().astype(str).head(top_n)
            if serie.empty:
                continue
            score = serie.apply(lambda x: self.normalizar_nombre_persona(self.homologar_nombre_grupo(x, nombres_grupo)) in grupo_norm).sum()
            if score > mejor_score:
                mejor_score = score
                mejor_col = col

        if mejor_score <= 0:
            return None
        return mejor_col

    def detectar_columna_fecha_por_nombre(self, df: pd.DataFrame, posibles: list) -> Optional[str]:
        return find_col(df, posibles)

    def preparar_base_preradicados_grupo_a(self, df_preradicados: pd.DataFrame, nombres_grupo_a: list) -> pd.DataFrame:
        df = self.normalizar_encabezados_sidcar(df_preradicados.copy())

        col_numero = find_col(df, ["Número", "Numero", "No. Doc", "No. Doc:", "Pre-Radicado", "Pre Radicado"])
        col_tipo = find_col(df, ["Tipo Documento", "Tipo"])
        col_estado = find_col(df, ["Estado"])
        col_asunto = find_col(df, ["Asunto", "Titulo", "Título"])
        col_fecha = find_col(df, ["Fecha", "Información de Creación", "Informacion de Creacion", "Fecha Creacion", "Fecha de Creación"])

        col_persona = find_col(df, ["Usuario", "Responsable", "Elaboró", "Elaboro", "Nombre", "Persona", "Creado Por", "Creado por"])
        if col_persona is None:
            col_persona = self.detectar_columna_persona_por_contenido(df, nombres_grupo_a)

        faltantes = {
            "numero_preradicado": col_numero,
            "tipo": col_tipo,
            "estado": col_estado,
            "asunto": col_asunto,
            "persona": col_persona,
        }
        faltan = [k for k, v in faltantes.items() if v is None]
        if faltan:
            raise KeyError(f"❌ En PRERADICADOS faltan columnas críticas: {faltan}. Columnas disponibles: {list(df.columns)}")

        df["nombre_base_preradicado"] = df[col_persona].astype(str)
        df["nombre_persona"] = df["nombre_base_preradicado"].apply(lambda x: self.homologar_nombre_grupo(x, nombres_grupo_a))

        grupo_norm = [self.normalizar_nombre_persona(x) for x in nombres_grupo_a]
        df = df[df["nombre_persona"].apply(lambda x: self.normalizar_nombre_persona(x) in grupo_norm)].copy()

        df["numero_preradicado"] = df[col_numero].astype(str).str.strip()
        df["tipo_documento_preradicado"] = df[col_tipo].astype(str).str.strip() if col_tipo else ""
        df["estado_preradicado"] = df[col_estado].astype(str).str.strip() if col_estado else ""
        df["asunto_preradicado"] = df[col_asunto].astype(str).str.strip() if col_asunto else ""
        df["fecha_preradicado"] = pd.to_datetime(df[col_fecha], errors="coerce") if col_fecha else pd.NaT

        col_expediente = find_col(df, ["Expediente", "Info SAE Expediente", "Numero Expediente"])
        if col_expediente:
            df["expediente"] = df[col_expediente].apply(self.extraer_expediente_texto)
        else:
            df["expediente"] = df["asunto_preradicado"].apply(self.extraer_expediente_texto)

        df = df[[
            "expediente",
            "nombre_persona",
            "nombre_base_preradicado",
            "numero_preradicado",
            "tipo_documento_preradicado",
            "estado_preradicado",
            "fecha_preradicado",
            "asunto_preradicado",
        ]].copy()

        df["expediente"] = df["expediente"].apply(self.limpiar_expediente)
        df = df[df["expediente"] != ""].copy()
        return df

    def preparar_base_radicados_grupo_a(self, df_radicados: pd.DataFrame, nombres_grupo_a: list) -> pd.DataFrame:
        df = self.normalizar_encabezados_sidcar(df_radicados.copy())

        col_radicado = find_col(df, ["# Radicado", "Radicado", "No. Radicado", "No Radicado"])
        col_tipo = find_col(df, ["Tipo"])
        col_fecha = find_col(df, ["F Radicado", "Fecha Radicado", "Fecha"])
        col_asunto = find_col(df, ["Asunto", "Titulo", "Título"])
        col_estado = find_col(df, ["Estado"])
        col_preradicado = find_col(df, ["Pre-Radicado", "Pre Radicado", "Preradicado"])

        col_persona = find_col(df, ["Elaboró", "Elaboro", "Responsable", "Usuario", "Nombre", "Persona"])
        if col_persona is None:
            col_persona = self.detectar_columna_persona_por_contenido(df, nombres_grupo_a)

        faltantes = {
            "radicado": col_radicado,
            "tipo": col_tipo,
            "fecha": col_fecha,
            "asunto": col_asunto,
            "estado": col_estado,
            "persona": col_persona,
        }
        faltan = [k for k, v in faltantes.items() if v is None]
        if faltan:
            raise KeyError(f"❌ En RADICADOS faltan columnas críticas: {faltan}. Columnas disponibles: {list(df.columns)}")

        df["nombre_base_radicado"] = df[col_persona].astype(str)
        df["nombre_persona"] = df["nombre_base_radicado"].apply(lambda x: self.homologar_nombre_grupo(x, nombres_grupo_a))

        grupo_norm = [self.normalizar_nombre_persona(x) for x in nombres_grupo_a]
        df = df[df["nombre_persona"].apply(lambda x: self.normalizar_nombre_persona(x) in grupo_norm)].copy()

        df["numero_radicado"] = df[col_radicado].astype(str).str.strip()
        df["tipo_documento_radicado"] = df[col_tipo].astype(str).str.strip() if col_tipo else ""
        df["estado_radicado"] = df[col_estado].astype(str).str.strip() if col_estado else ""
        df["asunto_radicado"] = df[col_asunto].astype(str).str.strip() if col_asunto else ""
        df["fecha_radicado"] = pd.to_datetime(df[col_fecha], errors="coerce") if col_fecha else pd.NaT
        df["numero_preradicado_ref"] = df[col_preradicado].astype(str).str.strip() if col_preradicado else ""

        col_expediente = find_col(df, ["Info SAE Expediente", "Expediente", "Numero Expediente"])
        if col_expediente:
            df["expediente"] = df[col_expediente].apply(self.extraer_expediente_texto)
        else:
            df["expediente"] = df["asunto_radicado"].apply(self.extraer_expediente_texto)

        df = df[[
            "expediente",
            "nombre_persona",
            "nombre_base_radicado",
            "numero_radicado",
            "tipo_documento_radicado",
            "estado_radicado",
            "fecha_radicado",
            "asunto_radicado",
            "numero_preradicado_ref",
        ]].copy()

        df["expediente"] = df["expediente"].apply(self.limpiar_expediente)
        df = df[df["expediente"] != ""].copy()
        return df

    def _join_unique(self, serie: pd.Series) -> str:
        vals = []
        for x in serie:
            if pd.isna(x):
                continue
            sx = str(x).strip()
            if sx and sx.lower() != "nan":
                vals.append(sx)
        return " | ".join(sorted(set(vals)))

    def consolidar_grupo_a_sidcar(self, df_preradicados: pd.DataFrame, df_radicados: pd.DataFrame, nombres_grupo_a: list) -> pd.DataFrame:
        pre = self.preparar_base_preradicados_grupo_a(df_preradicados, nombres_grupo_a)
        rad = self.preparar_base_radicados_grupo_a(df_radicados, nombres_grupo_a)

        pre_agg = pre.groupby(["expediente", "nombre_persona"], as_index=False).agg(
            nombre_en_preradicado=("nombre_base_preradicado", "first"),
            tiene_preradicado=("numero_preradicado", lambda s: any(str(x).strip() != "" for x in s)),
            numero_preradicado=("numero_preradicado", self._join_unique),
            tipo_documento_preradicado=("tipo_documento_preradicado", self._join_unique),
            estado_preradicado=("estado_preradicado", self._join_unique),
            fecha_preradicado=("fecha_preradicado", "min"),
            asunto_preradicado=("asunto_preradicado", "first"),
        )

        rad_agg = rad.groupby(["expediente", "nombre_persona"], as_index=False).agg(
            nombre_en_radicado=("nombre_base_radicado", "first"),
            tiene_radicado_firmado=("numero_radicado", lambda s: any(str(x).strip() != "" for x in s)),
            radicado=("numero_radicado", self._join_unique),
            tipo_radicado=("tipo_documento_radicado", self._join_unique),
            estado_radicado=("estado_radicado", self._join_unique),
            fecha_radicado=("fecha_radicado", "min"),
            asunto_radicado=("asunto_radicado", "first"),
            preradicado_relacionado=("numero_preradicado_ref", self._join_unique),
        )

        df_final = pre_agg.merge(rad_agg, on=["expediente", "nombre_persona"], how="outer")
        if "tiene_preradicado" not in df_final.columns:
            df_final["tiene_preradicado"] = False
        if "tiene_radicado_firmado" not in df_final.columns:
            df_final["tiene_radicado_firmado"] = False

        columnas_orden = [
            "expediente",
            "nombre_persona",
            "nombre_en_preradicado",
            "nombre_en_radicado",
            "tiene_preradicado",
            "numero_preradicado",
            "tipo_documento_preradicado",
            "estado_preradicado",
            "fecha_preradicado",
            "asunto_preradicado",
            "tiene_radicado_firmado",
            "radicado",
            "tipo_radicado",
            "estado_radicado",
            "fecha_radicado",
            "asunto_radicado",
            "preradicado_relacionado",
        ]

        for c in columnas_orden:
            if c not in df_final.columns:
                df_final[c] = ""

        df_final = df_final[columnas_orden].copy()
        df_final["fecha_preradicado"] = pd.to_datetime(df_final["fecha_preradicado"], errors="coerce")
        df_final["fecha_radicado"] = pd.to_datetime(df_final["fecha_radicado"], errors="coerce")

        df_final = df_final.sort_values(
            by=["nombre_persona", "expediente", "fecha_preradicado", "fecha_radicado"],
            ascending=[True, True, True, True],
        ).reset_index(drop=True)
        return df_final

    def preparar_base_documentos_elaborados_grupo_a(self, df_documentos_elaborados: pd.DataFrame, nombres_grupo_a: list) -> pd.DataFrame:
        df = self.normalizar_encabezados_sidcar(df_documentos_elaborados.copy())

        col_numero = find_col(df, [
            "Número", "Numero", "No. Doc", "No Doc", "Documento",
            "Consecutivo", "# Documento", "No. Documento", "Numero Documento",
        ])
        col_tipo = find_col(df, [
            "Tipo Documento", "Tipo", "Clase Documento", "Nombre Documento",
            "Documento Elaborado", "Tipo de Documento",
        ])
        col_estado = find_col(df, ["Estado", "Estado Documento", "Estado del Documento"])
        col_asunto = find_col(df, ["Asunto", "Titulo", "Título", "Descripcion", "Descripción", "Observacion", "Observación"])
        col_fecha = find_col(df, [
            "Fecha", "Información de Creación", "Informacion de Creacion",
            "Fecha Creacion", "Fecha de Creación", "Creado", "Fecha Elaboracion",
            "Fecha Elaboración",
        ])

        col_persona = find_col(df, ["Usuario", "Responsable", "Elaboró", "Elaboro", "Nombre", "Persona", "Creado Por", "Creado por"])
        if col_persona is None:
            col_persona = self.detectar_columna_persona_por_contenido(df, nombres_grupo_a)

        col_expediente = find_col(df, ["Expediente", "Info SAE Expediente", "Numero Expediente", "No Expediente"])

        faltantes = {"persona": col_persona}
        faltan = [k for k, v in faltantes.items() if v is None]
        if faltan:
            raise KeyError(f"❌ En DOCUMENTOS ELABORADOS faltan columnas críticas: {faltan}. Columnas disponibles: {list(df.columns)}")

        df["nombre_base_documentos"] = df[col_persona].astype(str)
        df["nombre_persona"] = df["nombre_base_documentos"].apply(lambda x: self.homologar_nombre_grupo(x, nombres_grupo_a))

        grupo_norm = [self.normalizar_nombre_persona(x) for x in nombres_grupo_a]
        df = df[df["nombre_persona"].apply(lambda x: self.normalizar_nombre_persona(x) in grupo_norm)].copy()

        df["numero_documento_elaborado"] = df[col_numero].astype(str).str.strip() if col_numero else ""
        df["tipo_documento_elaborado"] = df[col_tipo].astype(str).str.strip() if col_tipo else ""
        df["estado_documento_elaborado"] = df[col_estado].astype(str).str.strip() if col_estado else ""
        df["asunto_documento_elaborado"] = df[col_asunto].astype(str).str.strip() if col_asunto else ""
        df["fecha_documento_elaborado"] = pd.to_datetime(df[col_fecha], errors="coerce") if col_fecha else pd.NaT

        if col_expediente:
            df["expediente"] = df[col_expediente].apply(self.extraer_expediente_texto)
        else:
            df["expediente"] = df["asunto_documento_elaborado"].apply(self.extraer_expediente_texto)

        df = df[[
            "expediente",
            "nombre_persona",
            "nombre_base_documentos",
            "numero_documento_elaborado",
            "tipo_documento_elaborado",
            "estado_documento_elaborado",
            "fecha_documento_elaborado",
            "asunto_documento_elaborado",
        ]].copy()

        df["expediente"] = df["expediente"].apply(self.limpiar_expediente)
        df = df[df["expediente"] != ""].copy()
        return df

    def consolidar_documentos_elaborados_grupo_a(self, df_documentos_elaborados: pd.DataFrame, nombres_grupo_a: list) -> pd.DataFrame:
        docs = self.preparar_base_documentos_elaborados_grupo_a(df_documentos_elaborados, nombres_grupo_a)

        df_docs = docs.groupby(["expediente", "nombre_persona"], as_index=False).agg(
            nombre_en_documentos=("nombre_base_documentos", "first"),
            tiene_documento_elaborado=("numero_documento_elaborado", lambda s: any(str(x).strip() != "" for x in s)),
            numero_documento_elaborado=("numero_documento_elaborado", self._join_unique),
            tipo_documento_elaborado=("tipo_documento_elaborado", self._join_unique),
            estado_documento_elaborado=("estado_documento_elaborado", self._join_unique),
            fecha_documento_elaborado=("fecha_documento_elaborado", "min"),
            asunto_documento_elaborado=("asunto_documento_elaborado", "first"),
        )

        df_docs["fecha_documento_elaborado"] = pd.to_datetime(df_docs["fecha_documento_elaborado"], errors="coerce")
        df_docs = df_docs.sort_values(by=["nombre_persona", "expediente", "fecha_documento_elaborado"], ascending=[True, True, True]).reset_index(drop=True)
        return df_docs

    def consolidar_grupo_a_sidcar_y_documentos(
        self,
        df_preradicados: pd.DataFrame,
        df_radicados: pd.DataFrame,
        df_documentos_elaborados: pd.DataFrame,
        nombres_grupo_a: list,
    ):
        df_sidcar = self.consolidar_grupo_a_sidcar(df_preradicados=df_preradicados, df_radicados=df_radicados, nombres_grupo_a=nombres_grupo_a)
        df_docs = self.consolidar_documentos_elaborados_grupo_a(df_documentos_elaborados=df_documentos_elaborados, nombres_grupo_a=nombres_grupo_a)

        df_validacion = df_sidcar.merge(
            df_docs,
            on=["expediente", "nombre_persona"],
            how="outer",
            suffixes=("_sidcar", "_documentos"),
            indicator=True,
        )

        df_validacion["en_sidcar"] = df_validacion["_merge"].isin(["both", "left_only"])
        df_validacion["en_documentos"] = df_validacion["_merge"].isin(["both", "right_only"])
        df_validacion["coincide_expediente_persona"] = df_validacion["_merge"].eq("both")

        columnas_base = ["expediente", "nombre_persona", "coincide_expediente_persona", "en_sidcar", "en_documentos", "_merge"]
        columnas_resto = [c for c in df_validacion.columns if c not in columnas_base]
        df_validacion = df_validacion[columnas_base + columnas_resto].copy()
        df_validacion = df_validacion.sort_values(by=["nombre_persona", "expediente"], ascending=[True, True]).reset_index(drop=True)
        return df_sidcar, df_docs, df_validacion

    def exportar_grupo_a_sidcar(
        self,
        df_grupo_a: pd.DataFrame,
        df_documentos_elaborados: Optional[pd.DataFrame] = None,
        df_validacion: Optional[pd.DataFrame] = None,
        carpeta_salida: str = "output_SIDCAR",
        nombre_archivo: str = "reporte_grupo_A.xlsx",
    ) -> str:
        os.makedirs(carpeta_salida, exist_ok=True)
        ruta = os.path.join(carpeta_salida, nombre_archivo)

        with pd.ExcelWriter(ruta, engine="openpyxl") as writer:
            df_grupo_a.to_excel(writer, sheet_name="grupo_A_sidcar", index=False)
            if df_documentos_elaborados is not None:
                df_documentos_elaborados.to_excel(writer, sheet_name="grupo_A_documentos", index=False)
            if df_validacion is not None:
                df_validacion.to_excel(writer, sheet_name="grupo_A_validacion", index=False)

        print(f"✅ Archivo grupo A guardado en: {ruta}")
        return ruta
    
    def guardar_excel(self, df: pd.DataFrame, ruta_archivo: str, index: bool = False) -> str:
        ruta_archivo = str(ruta_archivo)
        carpeta = str(Path(ruta_archivo).parent)

        if carpeta and carpeta != ".":
            self.asegurar_carpeta(carpeta)

        df.to_excel(ruta_archivo, index=index)
        return ruta_archivo

    def guardar_varias_hojas(self, hojas: Dict[str, pd.DataFrame], ruta_archivo: str, index: bool = False) -> str:
        ruta_archivo = str(ruta_archivo)
        carpeta = str(Path(ruta_archivo).parent)

        if carpeta and carpeta != ".":
            self.asegurar_carpeta(carpeta)

        with pd.ExcelWriter(ruta_archivo, engine="openpyxl") as writer:
            for nombre_hoja, df_hoja in hojas.items():
                nombre_limpio = str(nombre_hoja).strip()[:31] if str(nombre_hoja).strip() else "Hoja1"
                df_hoja.to_excel(writer, sheet_name=nombre_limpio, index=index)

        return ruta_archivo

    # =========================
    # SEGUIMIENTO SAE + SIDCAR
    # BLOQUE ADITIVO - NO REEMPLAZA NADA EXISTENTE
    # =========================
    def _join_unique_seguimiento(self, serie: pd.Series) -> str:
        vals = []
        vistos = set()

        for x in serie:
            if pd.isna(x):
                continue

            sx = str(x).strip()
            if not sx or sx.lower() == "nan":
                continue

            key = sx.lower()
            if key not in vistos:
                vistos.add(key)
                vals.append(sx)

        return " | ".join(vals)

    def preparar_documentos_elaborados_seguimiento(
        self,
        df_documentos_elaborados: pd.DataFrame,
        nombres_objetivo: list
    ) -> pd.DataFrame:
        """
        Toma DocumentosElaborados como fuente SAE de asignación/gestión.
        No reemplaza la lógica previa de documentos elaborados.
        """
        df = self.normalizar_encabezados_sidcar(df_documentos_elaborados.copy())

        col_funcionario = find_col(df, ["Funcionario"])
        col_expediente = find_col(df, ["Expediente"])
        col_etapa = find_col(df, ["ACTIVIDAD ASIGNADA Etapa", "Etapa"])
        col_actividad = find_col(df, ["Actividad", "ACTIVIDAD ASIGNADA Actividad"])
        col_doc_tipo = find_col(df, ["ULTIMO DOCUMENTO ELABORADO Tipo", "Tipo"])
        col_doc_descripcion = find_col(df, ["Descripción", "Descripcion"])
        col_doc_fecha = find_col(df, ["Fecha"])
        col_doc_numero = find_col(df, ["Número", "Numero"])
        col_firmado_por = find_col(df, ["Firmado por"])

        faltantes = {
            "funcionario": col_funcionario,
            "expediente": col_expediente,
        }
        faltan = [k for k, v in faltantes.items() if v is None]

        if faltan:
            raise KeyError(
                f"❌ En DocumentosElaborados faltan columnas críticas: {faltan}. "
                f"Columnas disponibles: {list(df.columns)}"
            )

        df["funcionario_sae"] = df[col_funcionario].astype(str).apply(self.normalizar_nombre_persona)
        df["nombre_persona"] = df["funcionario_sae"].apply(
            lambda x: self.homologar_nombre_grupo(x, nombres_objetivo)
        )

        grupo_norm = [self.normalizar_nombre_persona(x) for x in nombres_objetivo]
        df = df[df["nombre_persona"].isin(grupo_norm)].copy()

        df["expediente"] = df[col_expediente].apply(self.limpiar_expediente)
        df["etapa_sae"] = df[col_etapa].astype(str).str.strip() if col_etapa else ""
        df["actividad_sae"] = df[col_actividad].astype(str).str.strip() if col_actividad else ""
        df["ultimo_documento_tipo_sae"] = df[col_doc_tipo].astype(str).str.strip() if col_doc_tipo else ""
        df["ultimo_documento_descripcion_sae"] = df[col_doc_descripcion].astype(str).str.strip() if col_doc_descripcion else ""
        df["fecha_ultimo_documento_sae"] = pd.to_datetime(df[col_doc_fecha], errors="coerce") if col_doc_fecha else pd.NaT
        df["numero_ultimo_documento_sae"] = df[col_doc_numero].astype(str).str.strip() if col_doc_numero else ""
        df["firmado_por_sae"] = df[col_firmado_por].astype(str).str.strip() if col_firmado_por else ""

        df = df[df["expediente"] != ""].copy()

        agg = df.groupby(["expediente", "nombre_persona"], as_index=False).agg(
            funcionario_sae=("funcionario_sae", "first"),
            etapa_sae=("etapa_sae", "last"),
            actividad_sae=("actividad_sae", "last"),
            ultimo_documento_tipo_sae=("ultimo_documento_tipo_sae", "last"),
            ultimo_documento_descripcion_sae=("ultimo_documento_descripcion_sae", "last"),
            fecha_ultimo_documento_sae=("fecha_ultimo_documento_sae", "max"),
            numero_ultimo_documento_sae=("numero_ultimo_documento_sae", "last"),
            firmado_por_sae=("firmado_por_sae", "last"),
        )

        return agg.sort_values(["nombre_persona", "expediente"]).reset_index(drop=True)

    def preparar_preradicados_seguimiento(
        self,
        df_preradicados: pd.DataFrame,
        nombres_objetivo: list
    ) -> pd.DataFrame:
        df = self.preparar_base_preradicados_grupo_a(
            df_preradicados=df_preradicados,
            nombres_grupo_a=nombres_objetivo
        ).copy()

        if "estado_preradicado" in df.columns:
            df["estado_preradicado_norm"] = df["estado_preradicado"].astype(str).str.strip().str.lower()
        else:
            df["estado_preradicado_norm"] = ""

        return df

    def preparar_radicados_seguimiento(
        self,
        df_radicados: pd.DataFrame,
        nombres_objetivo: list
    ) -> pd.DataFrame:
        df = self.preparar_base_radicados_grupo_a(
            df_radicados=df_radicados,
            nombres_grupo_a=nombres_objetivo
        ).copy()

        if "estado_radicado" in df.columns:
            df["estado_radicado_norm"] = df["estado_radicado"].astype(str).str.strip().str.lower()
        else:
            df["estado_radicado_norm"] = ""

        return df

    def consolidar_sidcar_seguimiento(
        self,
        df_preradicados: pd.DataFrame,
        df_radicados: pd.DataFrame,
        nombres_objetivo: list
    ) -> pd.DataFrame:
        pre = self.preparar_preradicados_seguimiento(df_preradicados, nombres_objetivo)
        rad = self.preparar_radicados_seguimiento(df_radicados, nombres_objetivo)

        pre_agg = pre.groupby(["expediente", "nombre_persona"], as_index=False).agg(
            tiene_preradicado=("numero_preradicado", lambda s: any(str(x).strip() != "" for x in s if not pd.isna(x))),
            numero_preradicado=("numero_preradicado", self._join_unique_seguimiento),
            tipo_documento_preradicado=("tipo_documento_preradicado", self._join_unique_seguimiento),
            estado_preradicado=("estado_preradicado", self._join_unique_seguimiento),
            fecha_preradicado=("fecha_preradicado", "max"),
            asunto_preradicado=("asunto_preradicado", "last"),
            nombre_base_preradicado=("nombre_base_preradicado", "first"),
        )

        rad_agg = rad.groupby(["expediente", "nombre_persona"], as_index=False).agg(
            tiene_radicado=("numero_radicado", lambda s: any(str(x).strip() != "" for x in s if not pd.isna(x))),
            numero_radicado=("numero_radicado", self._join_unique_seguimiento),
            tipo_documento_radicado=("tipo_documento_radicado", self._join_unique_seguimiento),
            estado_radicado=("estado_radicado", self._join_unique_seguimiento),
            fecha_radicado=("fecha_radicado", "max"),
            asunto_radicado=("asunto_radicado", "last"),
            nombre_base_radicado=("nombre_base_radicado", "first"),
            preradicado_relacionado=("numero_preradicado_ref", self._join_unique_seguimiento),
        )

        df = pre_agg.merge(rad_agg, on=["expediente", "nombre_persona"], how="outer")

        defaults = {
            "tiene_preradicado": False,
            "numero_preradicado": "",
            "tipo_documento_preradicado": "",
            "estado_preradicado": "",
            "fecha_preradicado": pd.NaT,
            "asunto_preradicado": "",
            "nombre_base_preradicado": "",
            "tiene_radicado": False,
            "numero_radicado": "",
            "tipo_documento_radicado": "",
            "estado_radicado": "",
            "fecha_radicado": pd.NaT,
            "asunto_radicado": "",
            "nombre_base_radicado": "",
            "preradicado_relacionado": "",
        }

        for col, default in defaults.items():
            if col not in df.columns:
                df[col] = default

        df["fecha_preradicado"] = pd.to_datetime(df["fecha_preradicado"], errors="coerce")
        df["fecha_radicado"] = pd.to_datetime(df["fecha_radicado"], errors="coerce")

        return df.sort_values(["nombre_persona", "expediente"]).reset_index(drop=True)

    def clasificar_estado_pelota_caliente_seguimiento(self, row: pd.Series) -> str:
        tiene_pre = bool(row.get("tiene_preradicado", False))
        tiene_rad = bool(row.get("tiene_radicado", False))

        estado_pre = str(row.get("estado_preradicado", "")).lower()
        asunto_pre = str(row.get("asunto_preradicado", "")).lower()
        estado_rad = str(row.get("estado_radicado", "")).lower()
        asunto_rad = str(row.get("asunto_radicado", "")).lower()

        texto_total = " | ".join([estado_pre, asunto_pre, estado_rad, asunto_rad])

        if tiene_rad:
            return "Firmado/Finalizado"

        if not tiene_pre:
            return "Sin iniciar"

        if any(x in texto_total for x in [
            "devuelto", "devolucion", "devolución", "observacion", "observación"
        ]):
            return "Devuelto"

        if any(x in texto_total for x in [
            "enviado para revision",
            "enviado para revisión",
            "enviado para aprobacion",
            "enviado para aprobación",
            "aprobado",
            "vobo",
            "vo.bo",
            "enviado para vobo",
            "enviado para vo.bo",
        ]):
            return "Aprobado VoBo"

        return "En creación"

    def definir_responsable_pelota_caliente_seguimiento(self, row: pd.Series) -> str:
        estado = str(row.get("estado_pelota_caliente", "")).strip()
        persona = str(row.get("nombre_persona", "")).strip()

        if estado in ["Sin iniciar", "En creación", "Devuelto"]:
            return persona

        if estado == "Aprobado VoBo":
            return f"REVISOR / {persona}"

        if estado == "Firmado/Finalizado":
            return "FINALIZADO"

        return persona

    def construir_reporte_seguimiento_sae_sidcar(
        self,
        df_documentos_elaborados: pd.DataFrame,
        df_preradicados: pd.DataFrame,
        df_radicados: pd.DataFrame,
        nombres_objetivo: list,
    ) -> pd.DataFrame:
        df_sae = self.preparar_documentos_elaborados_seguimiento(
            df_documentos_elaborados=df_documentos_elaborados,
            nombres_objetivo=nombres_objetivo,
        )

        df_sidcar = self.consolidar_sidcar_seguimiento(
            df_preradicados=df_preradicados,
            df_radicados=df_radicados,
            nombres_objetivo=nombres_objetivo,
        )

        df_final = df_sae.merge(
            df_sidcar,
            on=["expediente", "nombre_persona"],
            how="left"
        )

        defaults = {
            "tiene_preradicado": False,
            "numero_preradicado": "",
            "tipo_documento_preradicado": "",
            "estado_preradicado": "",
            "fecha_preradicado": pd.NaT,
            "asunto_preradicado": "",
            "tiene_radicado": False,
            "numero_radicado": "",
            "tipo_documento_radicado": "",
            "estado_radicado": "",
            "fecha_radicado": pd.NaT,
            "asunto_radicado": "",
            "preradicado_relacionado": "",
        }

        for col, default in defaults.items():
            if col not in df_final.columns:
                df_final[col] = default

        df_final["estado_pelota_caliente"] = df_final.apply(
            self.clasificar_estado_pelota_caliente_seguimiento,
            axis=1
        )
        df_final["responsable_pelota_caliente"] = df_final.apply(
            self.definir_responsable_pelota_caliente_seguimiento,
            axis=1
        )

        columnas_orden = [
            "expediente",
            "nombre_persona",
            "funcionario_sae",
            "etapa_sae",
            "actividad_sae",
            "ultimo_documento_tipo_sae",
            "ultimo_documento_descripcion_sae",
            "fecha_ultimo_documento_sae",
            "numero_ultimo_documento_sae",
            "firmado_por_sae",
            "tiene_preradicado",
            "numero_preradicado",
            "tipo_documento_preradicado",
            "estado_preradicado",
            "fecha_preradicado",
            "asunto_preradicado",
            "tiene_radicado",
            "numero_radicado",
            "tipo_documento_radicado",
            "estado_radicado",
            "fecha_radicado",
            "asunto_radicado",
            "preradicado_relacionado",
            "estado_pelota_caliente",
            "responsable_pelota_caliente",
        ]

        for c in columnas_orden:
            if c not in df_final.columns:
                df_final[c] = ""

        df_final = df_final[columnas_orden].copy()
        df_final = df_final.sort_values(
            by=["nombre_persona", "estado_pelota_caliente", "expediente"],
            ascending=[True, True, True]
        ).reset_index(drop=True)

        return df_final

    def construir_resumen_seguimiento_sae_sidcar(self, df_final: pd.DataFrame) -> pd.DataFrame:
        resumen = (
            df_final.groupby(["nombre_persona", "estado_pelota_caliente"], dropna=False)
            .agg(
                total_expedientes=("expediente", "nunique"),
                total_registros=("expediente", "count"),
            )
            .reset_index()
            .sort_values(["nombre_persona", "estado_pelota_caliente"])
        )
        return resumen

    def exportar_reporte_seguimiento_sae_sidcar(
        self,
        df_final: pd.DataFrame,
        carpeta_salida: str = "output_SIDCAR",
        nombre_archivo: str = "seguimiento_sae_sidcar.xlsx",
    ) -> str:
        self.asegurar_carpeta(carpeta_salida)
        ruta = os.path.join(carpeta_salida, nombre_archivo)

        resumen = self.construir_resumen_seguimiento_sae_sidcar(df_final)

        with pd.ExcelWriter(ruta, engine="openpyxl") as writer:
            df_final.to_excel(writer, sheet_name="seguimiento_detalle", index=False)
            resumen.to_excel(writer, sheet_name="seguimiento_resumen", index=False)

        print(f"✅ Reporte guardado en: {ruta}")
        return ruta



# ============================================================
# PARCHE ADITIVO 2026-03-26
# Ajuste de seguimiento SAE + SIDCAR sin eliminar lógica previa.
# Sobrescribe únicamente los métodos de seguimiento para:
# - usar OUTER MERGE real entre SAE, preradicados y radicados
# - conservar expedientes/personas presentes en cualquiera de las bases
# - calcular estados de "pelota caliente" sobre el consolidado final
# ============================================================

def _utils_join_unique_ordered(self, serie: pd.Series) -> str:
    vals = []
    seen = set()
    for x in serie:
        if pd.isna(x):
            continue
        sx = str(x).strip()
        if not sx or sx.lower() == "nan":
            continue
        key = sx.lower()
        if key not in seen:
            seen.add(key)
            vals.append(sx)
    return " | ".join(vals)


def _utils_texto_estado(self, *vals) -> str:
    partes = []
    for v in vals:
        if pd.isna(v) if hasattr(pd, 'isna') else False:
            continue
        s = str(v).strip()
        if s:
            partes.append(self.quitar_tildes(s).lower())
    return " | ".join(partes)


def _utils_homologar_persona_df(self, df: pd.DataFrame, col_persona: str, nombres_objetivo: list, col_salida_base: str, col_salida_norm: str = "nombre_persona") -> pd.DataFrame:
    out = df.copy()
    out[col_salida_base] = out[col_persona].astype(str)
    out[col_salida_norm] = out[col_salida_base].apply(lambda x: self.homologar_nombre_grupo(x, nombres_objetivo))
    grupo_norm = [self.normalizar_nombre_persona(x) for x in nombres_objetivo]
    out = out[out[col_salida_norm].apply(lambda x: self.normalizar_nombre_persona(x) in grupo_norm)].copy()
    return out


def _utils_preparar_documentos_elaborados_seguimiento(self, df_documentos_elaborados: pd.DataFrame, nombres_objetivo: list) -> pd.DataFrame:
    df = self.normalizar_encabezados_sidcar(df_documentos_elaborados.copy())

    col_funcionario = find_col(df, ["Funcionario", "Responsable", "Usuario", "Elaboró", "Elaboro"])
    col_expediente = find_col(df, ["Expediente", "Info SAE Expediente", "Numero Expediente"])
    col_etapa = find_col(df, ["ACTIVIDAD ASIGNADA Etapa", "Etapa"])
    col_actividad = find_col(df, ["ACTIVIDAD ASIGNADA Actividad", "Actividad"])
    col_doc_tipo = find_col(df, ["ULTIMO DOCUMENTO ELABORADO Tipo", "Tipo", "Tipo Documento"])
    col_doc_descripcion = find_col(df, ["ULTIMO DOCUMENTO ELABORADO Descripción", "ULTIMO DOCUMENTO ELABORADO Descripcion", "Descripción", "Descripcion", "Asunto"])
    col_doc_fecha = find_col(df, ["ULTIMO DOCUMENTO ELABORADO Fecha", "Fecha", "Fecha Elaboracion", "Fecha Elaboración"])
    col_doc_numero = find_col(df, ["ULTIMO DOCUMENTO ELABORADO Número", "ULTIMO DOCUMENTO ELABORADO Numero", "Número", "Numero", "Documento"])
    col_firmado_por = find_col(df, ["Firmado por", "Firmado Por"])

    faltantes = {"funcionario": col_funcionario, "expediente": col_expediente}
    faltan = [k for k, v in faltantes.items() if v is None]
    if faltan:
        raise KeyError(f"❌ En DocumentosElaborados faltan columnas críticas: {faltan}. Columnas disponibles: {list(df.columns)}")

    df = _utils_homologar_persona_df(self, df, col_funcionario, nombres_objetivo, "funcionario_sae")
    df["expediente"] = df[col_expediente].apply(self.extraer_expediente_texto if col_expediente else self.limpiar_expediente)
    df["expediente"] = df["expediente"].apply(self.limpiar_expediente)
    df["etapa_sae"] = df[col_etapa].astype(str).str.strip() if col_etapa else ""
    df["actividad_sae"] = df[col_actividad].astype(str).str.strip() if col_actividad else ""
    df["ultimo_documento_tipo_sae"] = df[col_doc_tipo].astype(str).str.strip() if col_doc_tipo else ""
    df["ultimo_documento_descripcion_sae"] = df[col_doc_descripcion].astype(str).str.strip() if col_doc_descripcion else ""
    df["fecha_ultimo_documento_sae"] = pd.to_datetime(df[col_doc_fecha], errors="coerce") if col_doc_fecha else pd.NaT
    df["numero_ultimo_documento_sae"] = df[col_doc_numero].astype(str).str.strip() if col_doc_numero else ""
    df["firmado_por_sae"] = df[col_firmado_por].astype(str).str.strip() if col_firmado_por else ""
    df["en_sae"] = True

    df = df[df["expediente"] != ""].copy()

    agg = df.groupby(["expediente", "nombre_persona"], as_index=False).agg(
        funcionario_sae=("funcionario_sae", "first"),
        etapa_sae=("etapa_sae", "last"),
        actividad_sae=("actividad_sae", "last"),
        ultimo_documento_tipo_sae=("ultimo_documento_tipo_sae", _utils_join_unique_ordered.__get__(self, type(self))),
        ultimo_documento_descripcion_sae=("ultimo_documento_descripcion_sae", "last"),
        fecha_ultimo_documento_sae=("fecha_ultimo_documento_sae", "max"),
        numero_ultimo_documento_sae=("numero_ultimo_documento_sae", _utils_join_unique_ordered.__get__(self, type(self))),
        firmado_por_sae=("firmado_por_sae", _utils_join_unique_ordered.__get__(self, type(self))),
        en_sae=("en_sae", "max"),
    )
    return agg.sort_values(["nombre_persona", "expediente"]).reset_index(drop=True)


def _utils_preparar_preradicados_seguimiento(self, df_preradicados: pd.DataFrame, nombres_objetivo: list) -> pd.DataFrame:
    df = self.normalizar_encabezados_sidcar(df_preradicados.copy())

    col_numero = find_col(df, ["Número", "Numero", "No. Doc", "No. Doc:", "Pre-Radicado", "Pre Radicado"])
    col_tipo = find_col(df, ["Tipo Documento", "Tipo"])
    col_estado = find_col(df, ["Estado", "Etapa"])
    col_asunto = find_col(df, ["Asunto", "Titulo", "Título", "Descripción", "Descripcion"])
    col_fecha = find_col(df, ["Fecha", "Información de Creación", "Informacion de Creacion", "Fecha Creacion", "Fecha de Creación"])
    col_persona = find_col(df, ["Usuario", "Responsable", "Elaboró", "Elaboro", "Nombre", "Persona", "Creado Por", "Creado por"])
    col_revisor = find_col(df, ["Revisó", "Reviso", "Revisor", "VoBo", "Vobo", "Aprobó", "Aprobo"])
    col_expediente = find_col(df, ["Expediente", "Info SAE Expediente", "Numero Expediente"])

    if col_persona is None:
        col_persona = self.detectar_columna_persona_por_contenido(df, nombres_objetivo)

    faltantes = {"numero_preradicado": col_numero, "estado": col_estado, "asunto": col_asunto, "persona": col_persona}
    faltan = [k for k, v in faltantes.items() if v is None]
    if faltan:
        raise KeyError(f"❌ En PRERADICADOS faltan columnas críticas: {faltan}. Columnas disponibles: {list(df.columns)}")

    df = _utils_homologar_persona_df(self, df, col_persona, nombres_objetivo, "nombre_base_preradicado")
    df["numero_preradicado"] = df[col_numero].astype(str).str.strip()
    df["tipo_documento_preradicado"] = df[col_tipo].astype(str).str.strip() if col_tipo else ""
    df["estado_preradicado"] = df[col_estado].astype(str).str.strip() if col_estado else ""
    df["asunto_preradicado"] = df[col_asunto].astype(str).str.strip() if col_asunto else ""
    df["fecha_preradicado"] = pd.to_datetime(df[col_fecha], errors="coerce") if col_fecha else pd.NaT
    df["revisor_pre"] = df[col_revisor].astype(str).str.strip() if col_revisor else ""
    if col_expediente:
        df["expediente"] = df[col_expediente].apply(self.extraer_expediente_texto)
    else:
        df["expediente"] = df["asunto_preradicado"].apply(self.extraer_expediente_texto)
    df["expediente"] = df["expediente"].apply(self.limpiar_expediente)
    df["en_preradicados"] = True
    df = df[df["expediente"] != ""].copy()
    return df[[
        "expediente", "nombre_persona", "nombre_base_preradicado", "numero_preradicado",
        "tipo_documento_preradicado", "estado_preradicado", "fecha_preradicado",
        "asunto_preradicado", "revisor_pre", "en_preradicados"
    ]].copy()


def _utils_preparar_radicados_seguimiento(self, df_radicados: pd.DataFrame, nombres_objetivo: list) -> pd.DataFrame:
    df = self.normalizar_encabezados_sidcar(df_radicados.copy())

    col_radicado = find_col(df, ["# Radicado", "Radicado", "No. Radicado", "No Radicado"])
    col_tipo = find_col(df, ["Tipo", "Tipo Documento"])
    col_fecha = find_col(df, ["F Radicado", "Fecha Radicado", "Fecha"])
    col_asunto = find_col(df, ["Asunto", "Titulo", "Título", "Descripción", "Descripcion"])
    col_estado = find_col(df, ["Estado"])
    col_preradicado = find_col(df, ["Pre-Radicado", "Pre Radicado", "Preradicado"])
    col_persona = find_col(df, ["Elaboró", "Elaboro", "Responsable", "Usuario", "Nombre", "Persona"])
    col_revisor = find_col(df, ["Revisó", "Reviso", "Revisor", "Aprobó", "Aprobo", "Firmó", "Firmo", "Radicó", "Radico"])
    col_expediente = find_col(df, ["Info SAE Expediente", "Expediente", "Numero Expediente"])

    if col_persona is None:
        col_persona = self.detectar_columna_persona_por_contenido(df, nombres_objetivo)

    faltantes = {"radicado": col_radicado, "estado": col_estado, "asunto": col_asunto, "persona": col_persona}
    faltan = [k for k, v in faltantes.items() if v is None]
    if faltan:
        raise KeyError(f"❌ En RADICADOS faltan columnas críticas: {faltan}. Columnas disponibles: {list(df.columns)}")

    df = _utils_homologar_persona_df(self, df, col_persona, nombres_objetivo, "nombre_base_radicado")
    df["numero_radicado"] = df[col_radicado].astype(str).str.strip()
    df["tipo_documento_radicado"] = df[col_tipo].astype(str).str.strip() if col_tipo else ""
    df["estado_radicado"] = df[col_estado].astype(str).str.strip() if col_estado else ""
    df["asunto_radicado"] = df[col_asunto].astype(str).str.strip() if col_asunto else ""
    df["fecha_radicado"] = pd.to_datetime(df[col_fecha], errors="coerce") if col_fecha else pd.NaT
    df["numero_preradicado_ref"] = df[col_preradicado].astype(str).str.strip() if col_preradicado else ""
    df["revisor_rad"] = df[col_revisor].astype(str).str.strip() if col_revisor else ""
    if col_expediente:
        df["expediente"] = df[col_expediente].apply(self.extraer_expediente_texto)
    else:
        df["expediente"] = df["asunto_radicado"].apply(self.extraer_expediente_texto)
    df["expediente"] = df["expediente"].apply(self.limpiar_expediente)
    df["en_radicados"] = True
    df = df[df["expediente"] != ""].copy()
    return df[[
        "expediente", "nombre_persona", "nombre_base_radicado", "numero_radicado",
        "tipo_documento_radicado", "estado_radicado", "fecha_radicado", "asunto_radicado",
        "numero_preradicado_ref", "revisor_rad", "en_radicados"
    ]].copy()


def _utils_consolidar_sidcar_seguimiento(self, df_preradicados: pd.DataFrame, df_radicados: pd.DataFrame, nombres_objetivo: list) -> pd.DataFrame:
    pre = _utils_preparar_preradicados_seguimiento(self, df_preradicados, nombres_objetivo)
    rad = _utils_preparar_radicados_seguimiento(self, df_radicados, nombres_objetivo)

    joiner = _utils_join_unique_ordered.__get__(self, type(self))

    pre_agg = pre.groupby(["expediente", "nombre_persona"], as_index=False).agg(
        nombre_base_preradicado=("nombre_base_preradicado", "first"),
        tiene_preradicado=("numero_preradicado", lambda s: any(str(x).strip() != "" for x in s if not pd.isna(x))),
        numero_preradicado=("numero_preradicado", joiner),
        tipo_documento_preradicado=("tipo_documento_preradicado", joiner),
        estado_preradicado=("estado_preradicado", joiner),
        fecha_preradicado=("fecha_preradicado", "max"),
        asunto_preradicado=("asunto_preradicado", "last"),
        revisor_pre=("revisor_pre", joiner),
        en_preradicados=("en_preradicados", "max"),
    )

    rad_agg = rad.groupby(["expediente", "nombre_persona"], as_index=False).agg(
        nombre_base_radicado=("nombre_base_radicado", "first"),
        tiene_radicado=("numero_radicado", lambda s: any(str(x).strip() != "" for x in s if not pd.isna(x))),
        numero_radicado=("numero_radicado", joiner),
        tipo_documento_radicado=("tipo_documento_radicado", joiner),
        estado_radicado=("estado_radicado", joiner),
        fecha_radicado=("fecha_radicado", "max"),
        asunto_radicado=("asunto_radicado", "last"),
        preradicado_relacionado=("numero_preradicado_ref", joiner),
        revisor_rad=("revisor_rad", joiner),
        en_radicados=("en_radicados", "max"),
    )

    df = pre_agg.merge(rad_agg, on=["expediente", "nombre_persona"], how="outer")
    defaults = {
        "nombre_base_preradicado": "", "tiene_preradicado": False, "numero_preradicado": "",
        "tipo_documento_preradicado": "", "estado_preradicado": "", "fecha_preradicado": pd.NaT,
        "asunto_preradicado": "", "revisor_pre": "", "en_preradicados": False,
        "nombre_base_radicado": "", "tiene_radicado": False, "numero_radicado": "",
        "tipo_documento_radicado": "", "estado_radicado": "", "fecha_radicado": pd.NaT,
        "asunto_radicado": "", "preradicado_relacionado": "", "revisor_rad": "", "en_radicados": False,
    }
    for col, default in defaults.items():
        if col not in df.columns:
            df[col] = default
    df["fecha_preradicado"] = pd.to_datetime(df["fecha_preradicado"], errors="coerce")
    df["fecha_radicado"] = pd.to_datetime(df["fecha_radicado"], errors="coerce")
    return df.sort_values(["nombre_persona", "expediente"]).reset_index(drop=True)


def _utils_clasificar_estado_pelota_caliente_seguimiento(self, row: pd.Series) -> str:
    tiene_pre = bool(row.get("tiene_preradicado", False))
    tiene_rad = bool(row.get("tiene_radicado", False))
    texto = _utils_texto_estado(
        self,
        row.get("estado_preradicado", ""), row.get("asunto_preradicado", ""),
        row.get("estado_radicado", ""), row.get("asunto_radicado", ""),
        row.get("actividad_sae", ""), row.get("ultimo_documento_descripcion_sae", ""),
        row.get("ultimo_documento_tipo_sae", "")
    )

    if tiene_rad:
        return "Firmado/Finalizado"
    if not tiene_pre:
        return "Sin iniciar"
    if any(x in texto for x in ["devuelto", "devolucion", "observacion", "correccion", "ajuste"]):
        return "Devuelto"
    if any(x in texto for x in [
        "aprobado", "aprobada", "vobo", "vo.bo", "visto bueno", "revision", "revisión",
        "para firma", "firma director", "firma regional", "firmar"
    ]):
        return "Aprobado VoBo"
    return "En creación"


def _utils_definir_responsable_pelota_caliente_seguimiento(self, row: pd.Series) -> str:
    estado = str(row.get("estado_pelota_caliente", "")).strip()
    persona = str(row.get("nombre_persona", "")).strip()
    revisor = str(row.get("revisor_asociado", "")).strip()
    if estado in ["Sin iniciar", "En creación", "Devuelto"]:
        return persona
    if estado == "Aprobado VoBo":
        return revisor if revisor else f"REVISOR / {persona}"
    if estado == "Firmado/Finalizado":
        return "FINALIZADO"
    return persona


def _utils_construir_reporte_seguimiento_sae_sidcar(self, df_documentos_elaborados: pd.DataFrame, df_preradicados: pd.DataFrame, df_radicados: pd.DataFrame, nombres_objetivo: list) -> pd.DataFrame:
    df_sae = _utils_preparar_documentos_elaborados_seguimiento(self, df_documentos_elaborados, nombres_objetivo)
    df_sidcar = _utils_consolidar_sidcar_seguimiento(self, df_preradicados, df_radicados, nombres_objetivo)

    # OUTER MERGE REAL: conserva lo que exista en cualquiera de las bases.
    df_final = df_sae.merge(df_sidcar, on=["expediente", "nombre_persona"], how="outer")

    defaults = {
        "funcionario_sae": "", "etapa_sae": "", "actividad_sae": "", "ultimo_documento_tipo_sae": "",
        "ultimo_documento_descripcion_sae": "", "fecha_ultimo_documento_sae": pd.NaT,
        "numero_ultimo_documento_sae": "", "firmado_por_sae": "", "en_sae": False,
        "nombre_base_preradicado": "", "tiene_preradicado": False, "numero_preradicado": "",
        "tipo_documento_preradicado": "", "estado_preradicado": "", "fecha_preradicado": pd.NaT,
        "asunto_preradicado": "", "revisor_pre": "", "en_preradicados": False,
        "nombre_base_radicado": "", "tiene_radicado": False, "numero_radicado": "",
        "tipo_documento_radicado": "", "estado_radicado": "", "fecha_radicado": pd.NaT,
        "asunto_radicado": "", "preradicado_relacionado": "", "revisor_rad": "", "en_radicados": False,
    }
    for col, default in defaults.items():
        if col not in df_final.columns:
            df_final[col] = default
        else:
            df_final[col] = df_final[col].fillna(default)

    # Si la persona viene solo desde SIDCAR, conservarla como funcionario SAE de referencia humana.
    df_final["funcionario_sae"] = df_final.apply(
        lambda r: r["funcionario_sae"] if str(r.get("funcionario_sae", "")).strip() else str(r.get("nombre_persona", "")).strip(),
        axis=1,
    )

    df_final["revisor_asociado"] = df_final.apply(
        lambda r: str(r.get("revisor_rad", "")).strip() or str(r.get("revisor_pre", "")).strip() or str(r.get("firmado_por_sae", "")).strip(),
        axis=1,
    )
    df_final["fuentes_presentes"] = df_final.apply(
        lambda r: " | ".join([
            fuente for fuente, ok in [
                ("SAE", bool(r.get("en_sae", False))),
                ("PRERADICADOS", bool(r.get("en_preradicados", False))),
                ("RADICADOS", bool(r.get("en_radicados", False))),
            ] if ok
        ]),
        axis=1,
    )

    df_final["estado_pelota_caliente"] = df_final.apply(lambda row: _utils_clasificar_estado_pelota_caliente_seguimiento(self, row), axis=1)
    df_final["responsable_pelota_caliente"] = df_final.apply(lambda row: _utils_definir_responsable_pelota_caliente_seguimiento(self, row), axis=1)

    columnas_orden = [
        "expediente", "nombre_persona", "funcionario_sae", "fuentes_presentes", "en_sae", "en_preradicados", "en_radicados",
        "etapa_sae", "actividad_sae", "ultimo_documento_tipo_sae", "ultimo_documento_descripcion_sae",
        "fecha_ultimo_documento_sae", "numero_ultimo_documento_sae", "firmado_por_sae",
        "tiene_preradicado", "numero_preradicado", "tipo_documento_preradicado", "estado_preradicado",
        "fecha_preradicado", "asunto_preradicado", "revisor_pre",
        "tiene_radicado", "numero_radicado", "tipo_documento_radicado", "estado_radicado",
        "fecha_radicado", "asunto_radicado", "preradicado_relacionado", "revisor_rad",
        "estado_pelota_caliente", "revisor_asociado", "responsable_pelota_caliente",
    ]
    for c in columnas_orden:
        if c not in df_final.columns:
            df_final[c] = ""

    df_final = df_final[columnas_orden].copy()
    df_final["fecha_ultimo_documento_sae"] = pd.to_datetime(df_final["fecha_ultimo_documento_sae"], errors="coerce")
    df_final["fecha_preradicado"] = pd.to_datetime(df_final["fecha_preradicado"], errors="coerce")
    df_final["fecha_radicado"] = pd.to_datetime(df_final["fecha_radicado"], errors="coerce")
    df_final = df_final.sort_values(
        by=["nombre_persona", "estado_pelota_caliente", "expediente"],
        ascending=[True, True, True],
    ).reset_index(drop=True)
    return df_final


def _utils_construir_resumen_seguimiento_sae_sidcar(self, df_final: pd.DataFrame) -> pd.DataFrame:
    resumen = (
        df_final.groupby(["nombre_persona", "estado_pelota_caliente"], dropna=False)
        .agg(
            total_expedientes=("expediente", "nunique"),
            total_registros=("expediente", "count"),
            con_sae=("en_sae", "sum"),
            con_preradicado=("en_preradicados", "sum"),
            con_radicado=("en_radicados", "sum"),
        )
        .reset_index()
        .sort_values(["nombre_persona", "estado_pelota_caliente"])
    )
    return resumen


def _utils_exportar_reporte_seguimiento_sae_sidcar(self, df_final: pd.DataFrame, carpeta_salida: str = "output_SIDCAR", nombre_archivo: str = "seguimiento_sae_sidcar.xlsx") -> str:
    self.asegurar_carpeta(carpeta_salida)
    ruta = os.path.join(carpeta_salida, nombre_archivo)
    resumen = _utils_construir_resumen_seguimiento_sae_sidcar(self, df_final)
    with pd.ExcelWriter(ruta, engine="openpyxl") as writer:
        df_final.to_excel(writer, sheet_name="seguimiento_detalle", index=False)
        resumen.to_excel(writer, sheet_name="seguimiento_resumen", index=False)
    print(f"✅ Reporte guardado en: {ruta}")
    return ruta


# Asignación explícita de métodos parcheados sobre Utils
Utils._join_unique_seguimiento = _utils_join_unique_ordered
Utils.preparar_documentos_elaborados_seguimiento = _utils_preparar_documentos_elaborados_seguimiento
Utils.preparar_preradicados_seguimiento = _utils_preparar_preradicados_seguimiento
Utils.preparar_radicados_seguimiento = _utils_preparar_radicados_seguimiento
Utils.consolidar_sidcar_seguimiento = _utils_consolidar_sidcar_seguimiento
Utils.clasificar_estado_pelota_caliente_seguimiento = _utils_clasificar_estado_pelota_caliente_seguimiento
Utils.definir_responsable_pelota_caliente_seguimiento = _utils_definir_responsable_pelota_caliente_seguimiento
Utils.construir_reporte_seguimiento_sae_sidcar = _utils_construir_reporte_seguimiento_sae_sidcar
Utils.construir_resumen_seguimiento_sae_sidcar = _utils_construir_resumen_seguimiento_sae_sidcar
Utils.exportar_reporte_seguimiento_sae_sidcar = _utils_exportar_reporte_seguimiento_sae_sidcar


# ============================================================
# PARCHE ADITIVO 2026-03-26 B
# Fechas de inicio / fin + días transcurridos + estado de vencimiento
# Mantiene todo lo existente y solo sobrescribe la lógica final
# del seguimiento SAE + SIDCAR.
# ============================================================

def _utils_calcular_fecha_inicio_expediente(self, row: pd.Series):
    """
    Fecha inicial del expediente por abogado.
    Prioridad:
    1) fecha_preradicado
    2) fecha_ultimo_documento_sae
    """
    fecha_pre = pd.to_datetime(row.get("fecha_preradicado"), errors="coerce")
    fecha_sae = pd.to_datetime(row.get("fecha_ultimo_documento_sae"), errors="coerce")

    if pd.notna(fecha_pre):
        return fecha_pre
    if pd.notna(fecha_sae):
        return fecha_sae
    return pd.NaT


def _utils_calcular_fecha_fin_expediente(self, row: pd.Series):
    """
    Fecha final del expediente.
    Regla:
    - si está finalizado / tiene radicado -> usar fecha_radicado
    - si por alguna razón está finalizado pero sin fecha_radicado, usar
      fecha_ultimo_documento_sae como respaldo operativo
    """
    estado = str(row.get("estado_pelota_caliente", "")).strip().lower()
    tiene_radicado = bool(row.get("tiene_radicado", False))

    fecha_rad = pd.to_datetime(row.get("fecha_radicado"), errors="coerce")
    fecha_sae = pd.to_datetime(row.get("fecha_ultimo_documento_sae"), errors="coerce")

    if tiene_radicado or estado == "firmado/finalizado":
        if pd.notna(fecha_rad):
            return fecha_rad
        if pd.notna(fecha_sae):
            return fecha_sae

    return pd.NaT


def _utils_calcular_dias_transcurridos_expediente(self, row: pd.Series):
    fecha_inicio = pd.to_datetime(row.get("fecha_inicio_expediente"), errors="coerce")
    fecha_fin = pd.to_datetime(row.get("fecha_fin_expediente"), errors="coerce")

    if pd.isna(fecha_inicio):
        return None

    fecha_corte = fecha_fin if pd.notna(fecha_fin) else pd.Timestamp.today().normalize()
    return int((fecha_corte.normalize() - fecha_inicio.normalize()).days)


def _utils_clasificar_estado_vencimiento(self, row: pd.Series) -> str:
    """
    Reglas:
    - Finalizado -> si ya terminó
    - Vencido -> > 30 días corridos
    - Por vencer -> desde 25 hasta 30 días corridos
    - En término -> < 25 días corridos
    - Sin fecha inicial -> no se puede medir
    """
    fecha_inicio = pd.to_datetime(row.get("fecha_inicio_expediente"), errors="coerce")
    fecha_fin = pd.to_datetime(row.get("fecha_fin_expediente"), errors="coerce")
    dias = row.get("dias_transcurridos_expediente")

    if pd.isna(fecha_inicio):
        return "Sin fecha inicial"

    if pd.notna(fecha_fin):
        return "Finalizado"

    if dias is None or pd.isna(dias):
        return "Sin fecha inicial"

    if dias > 30:
        return "Vencido"

    if 25 <= int(dias) <= 30:
        return "Por vencer"

    return "En término"


def _utils_construir_reporte_seguimiento_sae_sidcar_v2(self, df_documentos_elaborados: pd.DataFrame, df_preradicados: pd.DataFrame, df_radicados: pd.DataFrame, nombres_objetivo: list) -> pd.DataFrame:
    df_sae = _utils_preparar_documentos_elaborados_seguimiento(self, df_documentos_elaborados, nombres_objetivo)
    df_sidcar = _utils_consolidar_sidcar_seguimiento(self, df_preradicados, df_radicados, nombres_objetivo)

    # OUTER MERGE REAL: conserva lo que exista en cualquiera de las bases.
    df_final = df_sae.merge(df_sidcar, on=["expediente", "nombre_persona"], how="outer")

    defaults = {
        "funcionario_sae": "", "etapa_sae": "", "actividad_sae": "", "ultimo_documento_tipo_sae": "",
        "ultimo_documento_descripcion_sae": "", "fecha_ultimo_documento_sae": pd.NaT,
        "numero_ultimo_documento_sae": "", "firmado_por_sae": "", "en_sae": False,
        "nombre_base_preradicado": "", "tiene_preradicado": False, "numero_preradicado": "",
        "tipo_documento_preradicado": "", "estado_preradicado": "", "fecha_preradicado": pd.NaT,
        "asunto_preradicado": "", "revisor_pre": "", "en_preradicados": False,
        "nombre_base_radicado": "", "tiene_radicado": False, "numero_radicado": "",
        "tipo_documento_radicado": "", "estado_radicado": "", "fecha_radicado": pd.NaT,
        "asunto_radicado": "", "preradicado_relacionado": "", "revisor_rad": "", "en_radicados": False,
    }
    for col, default in defaults.items():
        if col not in df_final.columns:
            df_final[col] = default
        else:
            try:
                df_final[col] = df_final[col].fillna(default)
            except Exception:
                pass

    df_final["funcionario_sae"] = df_final.apply(
        lambda r: r["funcionario_sae"] if str(r.get("funcionario_sae", "")).strip() else str(r.get("nombre_persona", "")).strip(),
        axis=1,
    )

    df_final["revisor_asociado"] = df_final.apply(
        lambda r: str(r.get("revisor_rad", "")).strip() or str(r.get("revisor_pre", "")).strip() or str(r.get("firmado_por_sae", "")).strip(),
        axis=1,
    )

    df_final["fuentes_presentes"] = df_final.apply(
        lambda r: " | ".join([
            fuente for fuente, ok in [
                ("SAE", bool(r.get("en_sae", False))),
                ("PRERADICADOS", bool(r.get("en_preradicados", False))),
                ("RADICADOS", bool(r.get("en_radicados", False))),
            ] if ok
        ]),
        axis=1,
    )

    df_final["estado_pelota_caliente"] = df_final.apply(
        lambda row: _utils_clasificar_estado_pelota_caliente_seguimiento(self, row),
        axis=1
    )

    df_final["responsable_pelota_caliente"] = df_final.apply(
        lambda row: _utils_definir_responsable_pelota_caliente_seguimiento(self, row),
        axis=1
    )

    # Nuevas fechas y estados de control
    df_final["fecha_inicio_expediente"] = df_final.apply(
        lambda row: _utils_calcular_fecha_inicio_expediente(self, row),
        axis=1
    )

    df_final["fecha_fin_expediente"] = df_final.apply(
        lambda row: _utils_calcular_fecha_fin_expediente(self, row),
        axis=1
    )

    df_final["dias_transcurridos_expediente"] = df_final.apply(
        lambda row: _utils_calcular_dias_transcurridos_expediente(self, row),
        axis=1
    )

    df_final["estado_vencimiento"] = df_final.apply(
        lambda row: _utils_clasificar_estado_vencimiento(self, row),
        axis=1
    )

    columnas_orden = [
        "expediente", "nombre_persona", "funcionario_sae",
        "fuentes_presentes", "en_sae", "en_preradicados", "en_radicados",
        "etapa_sae", "actividad_sae", "ultimo_documento_tipo_sae", "ultimo_documento_descripcion_sae",
        "fecha_ultimo_documento_sae", "numero_ultimo_documento_sae", "firmado_por_sae",
        "tiene_preradicado", "numero_preradicado", "tipo_documento_preradicado", "estado_preradicado",
        "fecha_preradicado", "asunto_preradicado", "revisor_pre",
        "tiene_radicado", "numero_radicado", "tipo_documento_radicado", "estado_radicado",
        "fecha_radicado", "asunto_radicado", "preradicado_relacionado", "revisor_rad",
        "fecha_inicio_expediente", "fecha_fin_expediente", "dias_transcurridos_expediente", "estado_vencimiento",
        "estado_pelota_caliente", "revisor_asociado", "responsable_pelota_caliente",
    ]

    for c in columnas_orden:
        if c not in df_final.columns:
            df_final[c] = ""

    df_final = df_final[columnas_orden].copy()
    for c in ["fecha_ultimo_documento_sae", "fecha_preradicado", "fecha_radicado", "fecha_inicio_expediente", "fecha_fin_expediente"]:
        df_final[c] = pd.to_datetime(df_final[c], errors="coerce")

    df_final = df_final.sort_values(
        by=["nombre_persona", "estado_vencimiento", "estado_pelota_caliente", "expediente"],
        ascending=[True, True, True, True],
    ).reset_index(drop=True)

    return df_final


# Asignación explícita de métodos parcheados sobre Utils
Utils.calcular_fecha_inicio_expediente = _utils_calcular_fecha_inicio_expediente
Utils.calcular_fecha_fin_expediente = _utils_calcular_fecha_fin_expediente
Utils.calcular_dias_transcurridos_expediente = _utils_calcular_dias_transcurridos_expediente
Utils.clasificar_estado_vencimiento = _utils_clasificar_estado_vencimiento
Utils.construir_reporte_seguimiento_sae_sidcar = _utils_construir_reporte_seguimiento_sae_sidcar_v2
